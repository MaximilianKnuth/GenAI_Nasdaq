from docx import Document

def find_table_summary(doc, table_name):
    """
    Iterate over all non-empty paragraphs in the document and return a tuple (position, text) 
    for the first paragraph that contains the table name (case-insensitive).
    
    Parameters:
      doc (Document): A python-docx Document object.
      table_name (str): The table name to search for (e.g., "eqr").
      
    Returns:
      tuple: (position, paragraph_text) where position is the count (1-indexed) of non-empty paragraphs,
             and paragraph_text is the text of the found paragraph.
             If not found, returns (None, default_message).
    """
    count = 0
    for para in doc.paragraphs:
        if para.text.strip():  # if it's a non-empty paragraph
            count += 1
        if table_name.lower() in para.text.lower():
            return count, para.text.strip()
    return None, f"No paragraph containing '{table_name}' found."

def get_table_data_by_datatype(doc, table_index):
    """
    From the table at the given 1-indexed position, return a dictionary with:
      - Keys for each data type found, where each value is a list of column names of that data type.
      - A key "primary_key" whose value is a list of column names marked as primary keys.
    
    The table must have headers: "Column", "Data Type", and "Primary Key".
    
    Parameters:
      doc (Document): A python-docx Document object.
      table_index (int): The 1-indexed table number to process.
      
    Returns:
      dict: A dictionary grouping column names by their data type and including a "primary_key" key.
      
    Raises:
      ValueError: If the table index is invalid or the required headers are missing.
    """
    num_tables = len(doc.tables)
    if table_index < 1 or table_index > num_tables:
        raise ValueError(f"Invalid table index: {table_index}. Document has {num_tables} table(s).")
    
    table = doc.tables[table_index - 1]
    header = [cell.text.strip() for cell in table.rows[0].cells]
    # Create a lookup for header positions (using lowercase keys)
    header_map = {h.lower(): i for i, h in enumerate(header)}
    
    # Verify required headers exist.
    for required in ["column", "data type", "primary key"]:
        if required not in header_map:
            raise ValueError(f"The table does not contain the required header: '{required}'")
    
    col_idx = header_map["column"]
    dt_idx = header_map["data type"]
    pk_idx = header_map["primary key"]
    
    datatype_dict = {}
    primary_keys = []
    
    # Process each data row (skip header)
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        # Safety check: ensure we have enough cells
        if len(cells) <= max(col_idx, dt_idx, pk_idx):
            continue
        
        col_name = cells[col_idx]
        data_type = cells[dt_idx]
        pk_flag = cells[pk_idx]
        
        # Group column names by data type
        if data_type in datatype_dict:
            datatype_dict[data_type].append(col_name)
        else:
            datatype_dict[data_type] = [col_name]
        
        # Check for primary key marker "✅"
        if pk_flag == "✅":
            primary_keys.append(col_name)
    
    # Add primary key info as a separate key in the resulting dictionary.
    datatype_dict["primary_key"] = primary_keys
    
    return datatype_dict

def trunk_table_execute(table_name, doc_path="Data/Data Product Samples.docx"):
    """
    Open the DOCX file at the provided path, search for the paragraph containing the table name,
    and then use that paragraph's position (assumed to match the table index) to extract the table data.
    
    Parameters:
      table_name (str): The table name to search for (e.g., "eqr").
      doc_path (str): Path to the DOCX file.
    
    Returns:
      tuple: (summary, table_data)
        summary: the text from the paragraph containing the table name.
        table_data: a dictionary grouping column names by datatype and including primary keys.
    """
    doc = Document(doc_path)
    position, summary = find_table_summary(doc, table_name)
    if position is None:
        raise ValueError(f"Cannot find a paragraph containing '{table_name}' in the document.")
    
    # Use the found paragraph count as the table index.
    table_data = get_table_data_by_datatype(doc, position)
    return summary, table_data
